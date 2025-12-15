from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny   
from django.conf import settings
import re
from PyPDF2 import PdfReader
import os
from .analyzer import TenderBidAnalyzer

# tender_analyzer/views.py

class AnalyzePDFView(APIView):
    permission_classes = [AllowAny]  # या IsAuthenticated

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({"error": "No file uploaded"}, status=400)

        uploaded_file = request.FILES['file']

        # Save temporarily in media/temp/
        temp_dir = os.path.join(settings.MEDIA_ROOT, "temp")
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, uploaded_file.name)

        with open(temp_path, 'wb+') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)

        # अब analyze करो
        analyzer = TenderBidAnalyzer()
        result = analyzer.analyze_and_save(
            temp_path,
            output_dir=os.path.join(settings.MEDIA_ROOT, "analyzed")
        )

        # Temp file delete कर दो (optional)
        try:
            os.remove(temp_path)
        except:
            pass

        return Response(result, status=200)
    
    
    def safe(self, val):
        return val.strip() if isinstance(val, str) else ""

    def analyze_and_save(self, pdf_path, output_dir):
        try:
            reader = PdfReader(pdf_path)
            full_text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text

            links = re.findall(r'(https?://[^\s]+)', full_text)
            links = [self.safe(url).replace(")", "").replace("]", "") for url in links]

            document_links = [url for url in links if "documentdownload" in url]

            titles = [
                "tender notification", "annexre B", "annexure A",
                "format", "technical spec", "enclosure", "enclosure"
            ]

            document_objects = []
            for i, link in enumerate(document_links):
                document_objects.append({
                    "title": titles[i] if i < len(titles) else f"Document {i+1}",
                    "url": link
                })

            return {
                "status": True,
                "message": "Documents extracted successfully",
                "documents": document_objects
            }

        except Exception as e:
            return {
                "status": False,
                "error": str(e),
                "documents": []
            }
def extract_links_from_pdf(self, pdf_path):
    import fitz  # PyMuPDF
    links = []

    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            for link in page.get_links():
                if "uri" in link:
                    links.append(link["uri"])
        doc.close()
    except Exception as e:
        print("PDF Link Extract Error:", e)

    return links

            
            
            



from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import os
from django.conf import settings

class UploadBidPDFView(APIView):
    permission_classes = [AllowAny]  # या IsAuthenticated

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({"error": "No file uploaded"}, status=400)

        uploaded_file = request.FILES['file']
        filename = uploaded_file.name

        # Django के media/bids/ में save करो
        save_dir = os.path.join(settings.MEDIA_ROOT, "bids")
        os.makedirs(save_dir, exist_ok=True)
        file_path = os.path.join(save_dir, filename)

        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        # Frontend को relative path return करो
        relative_path = f"bids/{filename}"

        return Response({
            "message": "File uploaded successfully",
            "file_path": relative_path   # ← यही analyze में भेजना है
        }, status=201)
        
        
        
        
        
        
        
        
# tender_analyzer/views.py

import requests
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from rest_framework.views import APIView
import os
from django.conf import settings

@method_decorator(csrf_exempt, name='dispatch')
class DownloadSubdocsView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        data = request.data
        main_bid_filename = data.get("main_bid_filename")  # jaise "bc16110ae1.pdf"
        document_links = data.get("document_links", [])    # jo tune extract kiye hain

        if not main_bid_filename or not document_links:
            return Response({"error": "Missing data"}, status=400)

        # Subdocs folder banao
        subdoc_dir = os.path.join(settings.MEDIA_ROOT, "subdocs")
        os.makedirs(subdoc_dir, exist_ok=True)

        saved_files = []
        bid_prefix = os.path.splitext(main_bid_filename)[0]  # bc16110ae1

        for idx, doc in enumerate(document_links):
            url = doc.get("url")
            title = doc.get("title", f"Document_{idx+1}")

            try:
                print(f"Downloading → {title}: {url}")
                response = requests.get(url, timeout=40)

                if response.status_code != 200:
                    saved_files.append({"title": title, "error": "Download failed", "status_code": response.status_code})
                    continue

                # Safe filename
                ext = url.split(".")[-1].split("?")[0]
                if len(ext) > 5 or ext not in ["pdf", "doc", "docx", "xls", "xlsx", "zip"]:
                    ext = "pdf"

                filename = f"{bid_prefix}_document_{idx+1}.{ext}"
                filepath = os.path.join(subdoc_dir, filename)

                with open(filepath, "wb") as f:
                    f.write(response.content)

                local_url = f"/media/subdocs/{filename}"

                saved_files.append({
                    "title": title,
                    "original_url": url,
                    "local_url": local_url,
                    "filename": filename,
                    "size_kb": round(len(response.content)/1024, 2)
                })

            except Exception as e:
                saved_files.append({"title": title, "error": str(e)})

        return Response({
            "message": "Sub documents downloaded",
            "count": len(saved_files),
            "subdocs": saved_files
        }, status=200)
        
        
      
      
      
      
      
      
      
      
        
        
        
        
import os
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.response import Response
from app.models import Company
from .Bidparticipation import AutoDocService
from .models import GeneratedDocument
import uuid
from django.utils import timezone


class AutoDocGenerateAPI(APIView):
    permission_classes = []   # Testing only

    # ------------------------------------------
    # Helper: Convert AI text → Dictionary
    # ------------------------------------------
    def split_into_sections(self, ai_text):
        sections = {}
        current_title = None
        current_body = []

        for line in ai_text.split("\n"):
            line = line.strip()

            if line.startswith("### SECTION:"):
                if current_title:
                    sections[current_title] = "\n".join(current_body)

                current_title = line.replace("### SECTION:", "").strip()
                current_body = []

            else:
                current_body.append(line)

        if current_title:
            sections[current_title] = "\n".join(current_body)

        return sections

    # ------------------------------------------
    # POST API
    # ------------------------------------------
    def post(self, request):
        try:
            file = request.FILES.get("file")
            company_id = request.POST.get("company")
            custom_prompt = request.POST.get("prompt", "")

            if not file:
                return Response({"error": "File is required"}, status=400)
            if not company_id:
                return Response({"error": "Company is required"}, status=400)

            company = Company.objects.get(id=company_id)

            # 🔥 AI RAW TEXT
            ai_output = AutoDocService.process_document(file, company, custom_prompt)

            if ai_output.startswith("AI ERROR"):
                return Response({"error": ai_output}, status=500)

            # 🔥 Convert to dict
            sections = self.split_into_sections(ai_output)

            # --------------------------
            # HTML PREVIEW
            # --------------------------
            html_preview = ""
            for title, body in sections.items():
                html_preview += f"<h2>{title}</h2><p>{body.replace(chr(10), '<br>')}</p><hr>"

            # --------------------------
            # PDF GENERATION
            # --------------------------
            pdf_buffer = BytesIO()
            styles = getSampleStyleSheet()
            story = []

            for title, body in sections.items():
                story.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
                for line in body.split("\n"):
                    story.append(Paragraph(line, styles["Normal"]))

            pdf = SimpleDocTemplate(pdf_buffer)
            pdf.build(story)

            pdf_path = f"documentprepared/tender_{company_id}.pdf"
            default_storage.save(pdf_path, ContentFile(pdf_buffer.getvalue()))

            # --------------------------
            # DOCX GENERATION
            # --------------------------
            doc = Document()
            for title, body in sections.items():
                doc.add_heading(title, level=1)
                for line in body.split("\n"):
                    doc.add_paragraph(line)

            docx_buffer = BytesIO()
            doc.save(docx_buffer)

            docx_path = f"documentprepared/tender_{company_id}.docx"
            default_storage.save(docx_path, ContentFile(docx_buffer.getvalue()))

            # --------------------------
            # RESPONSE
            # --------------------------
            return Response({
                "success": True,
                "message": "Document generated successfully",
                "preview_html": html_preview,
                "pdf_url": settings.MEDIA_URL + pdf_path,
                "docx_url": settings.MEDIA_URL + docx_path
            })

        except Exception as e:
            return Response({"error": str(e)}, status=500)








# views.py mein kahin bhi daal de (bottom pe best)

from django.http import JsonResponse
from django.conf import settings
import os
from datetime import datetime

def list_prepared_documents(request):
    # JWT token check nahi karna padega kyunki DRF automatically karega agar tu IsAuthenticated laga de
    folder = os.path.join(settings.MEDIA_ROOT, "documentprepared")
    docs = []

    if not os.path.exists(folder):
        return JsonResponse({"documents": []})

    for file in sorted(os.listdir(folder), reverse=True):
        if file.endswith((".pdf", ".docx")):
            file_path = os.path.join(folder, file)
            timestamp = os.path.getctime(file_path)
            date_str = datetime.fromtimestamp(timestamp).strftime("%d %b %Y, %I:%M %p")

            docs.append({
                "name": file.replace(".pdf", "").replace(".docx", "").replace("_", " "),
                "pdf_url": f"/media/documentprepared/{file}" if file.endswith(".pdf") else None,
                "docx_url": f"/media/documentprepared/{file}" if file.endswith(".docx") else None,
                "date": date_str,
                "size_kb": round(os.path.getsize(file_path) / 1024, 1)
            })

    return JsonResponse({"documents": docs})