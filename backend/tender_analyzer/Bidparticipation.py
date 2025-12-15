# from openai import OpenAI
# from django.conf import settings
# from PyPDF2 import PdfReader
# import docx
# import os

# client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# class AutoDocService:

#     # 🎯 Read PDF
#     @staticmethod
#     def extract_pdf_text(file):
#         try:
#             reader = PdfReader(file)
#             text = ""
#             for page in reader.pages:
#                 txt = page.extract_text()
#                 if txt:
#                     text += txt + "\n"
#             return text
#         except:
#             return ""

#     # 🎯 Read DOCX
#     @staticmethod
#     def extract_docx_text(file):
#         try:
#             doc = docx.Document(file)
#             return "\n".join([p.text for p in doc.paragraphs])
#         except:
#             return ""

#     # 🎯 Auto detect file
#     @staticmethod
#     def read_uploaded_file(uploaded_file):
#         name = uploaded_file.name.lower()

#         if name.endswith(".pdf"):
#             return AutoDocService.extract_pdf_text(uploaded_file)

#         if name.endswith(".docx"):
#             return AutoDocService.extract_docx_text(uploaded_file)

#         try:
#             return uploaded_file.read().decode("utf-8", errors="ignore")
#         except:
#             return ""

#     # 🎯 BUILD AI PROMPT
#     @staticmethod
#     def build_prompt(company, bid_text, custom_prompt):

#         directors = "\n".join([
#             f"- {d.name}, Email: {d.email}, Phone: {d.phone}"
#             for d in company.directors.all()
#         ]) or "No directors found"

#         owners = "\n".join([
#             f"- {o.name}, Email: {o.email}, Phone: {o.phone}"
#             for o in company.owners.all()
#         ]) or "No owners found"

#         partners = "\n".join([
#             f"- {p.name}, Email: {p.email}, Phone: {p.phone}"
#             for p in company.partners.all()
#         ]) or "No partners found"

#         members = "\n".join([
#             f"- {m.name}, Email: {m.email}, Phone: {m.phone}"
#             for m in company.members.all()
#         ]) or "No members found"

#         prompt = f"""
# You are an expert AI document generator for Government tenders and GEM bids.

# ## COMPANY INFORMATION
# Company Name: {company.company_name}
# Company Type: {company.company_type}
# Major Activity: {company.major_activity}
# Nature of Business: {company.nature_of_business}

# Address:
# {company.company_address}

# GSTIN No: {company.gstin_no or "Not Available"}
# MSME No: {company.msme_no or "Not Available"}
# Enterprise Type: {company.enterprise_type or "Not Available"}

# Bank Details:
# - Account Holder: {company.account_holder_name or "N/A"}
# - Bank Name: {company.bank_name or "N/A"}
# - Account Number: {company.account_number or "N/A"}
# - IFSC: {company.ifsc_code or "N/A"}
# - Bank Phone: {company.bank_phone or "N/A"}

# ### Directors:
# {directors}

# ### Owners:
# {owners}

# ### Partners:
# {partners}

# ### Members:
# {members}

# ---------------------------------------------

# ## EXTRACTED TENDER/BID CONTENT:
# {bid_text}

# ---------------------------------------------

# ## USER CUSTOM INSTRUCTIONS:
# {custom_prompt}

# ---------------------------------------------

# ### TASK FOR AI:
# Using all the above information:
# Prepare all required tender documents including:
# - Cover letter  
# - Undertaking  
# - Declaration  
# - Eligibility confirmation  
# - Compliance statements  
# - Experience/turnover details (if applicable)  
# - Any tender-specific mandatory formats  

# Make the output clean, professional, well-structured, and tender-ready.
# Return only the final document text.
# """
#         return prompt

#     # 🎯 AI Request — FIXED FOR OPENAI V1+
#     @staticmethod
#     def ask_ai(prompt):
#         try:
#             response = client.chat.completions.create(
#                 model="gpt-4o-mini",
#                 messages=[
#                     {"role": "system", "content": "You generate professional tender documents."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 temperature=0.2
#             )

#             return response.choices[0].message.content

#         except Exception as e:
#             return f"AI ERROR: {str(e)}"

#     # 🎯 Main process
#     @staticmethod
#     def process_document(uploaded_file, company, custom_prompt):
#         bid_text = AutoDocService.read_uploaded_file(uploaded_file)
#         prompt = AutoDocService.build_prompt(company, bid_text, custom_prompt)
#         return AutoDocService.ask_ai(prompt)








import os
from PyPDF2 import PdfReader
import docx
from openai import OpenAI

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


class AutoDocService:

    # -------------------------------
    # PDF Extract
    # -------------------------------
    @staticmethod
    def extract_pdf_text(file):
        try:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except Exception:
            return ""

    # -------------------------------
    # DOCX Extract
    # -------------------------------
    @staticmethod
    def extract_docx_text(file):
        try:
            docf = docx.Document(file)
            return "\n".join([p.text for p in docf.paragraphs])
        except Exception:
            return ""

    # -------------------------------
    # Detect File Type
    # -------------------------------
    @staticmethod
    def read_uploaded_file(uploaded_file):
        name = uploaded_file.name.lower()

        if name.endswith(".pdf"):
            return AutoDocService.extract_pdf_text(uploaded_file)

        if name.endswith(".docx"):
            return AutoDocService.extract_docx_text(uploaded_file)

        # fallback → TXT
        try:
            return uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception:
            return ""

    # -------------------------------
    # Build Prompt
    # -------------------------------
    @staticmethod
    def build_prompt(company, bid_text, custom_prompt):

        directors = "\n".join([
            f"- {d.name} ({d.email}, {d.phone})"
            for d in company.directors.all()
        ]) or "N/A"

        owners = "\n".join([
            f"- {o.name} ({o.email}, {o.phone})"
            for o in company.owners.all()
        ]) or "N/A"

        partners = "\n".join([
            f"- {p.name} ({p.email}, {p.phone})"
            for p in company.partners.all()
        ]) or "N/A"

        members = "\n".join([
            f"- {m.name} ({m.email}, {m.phone})"
            for m in company.members.all()
        ]) or "N/A"

        return f"""
You are a Govt tender document generation expert AI.

### COMPANY INFORMATION
Name: {company.company_name}
Type: {company.company_type}
Major Activity: {company.major_activity}
Nature of Business: {company.nature_of_business}

Address:
{company.company_address}

GSTIN: {company.gstin_no or "N/A"}
MSME No: {company.msme_no or "N/A"}
Enterprise Type: {company.enterprise_type or "N/A"}

Contact Email: {company.contact_email or "N/A"}
Contact Phone: {company.contact_phone or "N/A"}
Designation: {company.designation or "Authorized Signatory"}
Experience: {company.years_of_experience or "3+ Years"}
Turnover: {company.annual_turnover or "As per MSME norms"}

### DIRECTORS
{directors}

### OWNERS
{owners}

### PARTNERS
{partners}

### MEMBERS
{members}

-----------------------------------------

### TENDER EXTRACTED TEXT
{bid_text}

-----------------------------------------

### CUSTOM USER INSTRUCTIONS
{custom_prompt}

-----------------------------------------

### TASK:
Generate the following sections EXACTLY in clean formatted text:

### SECTION: COVER LETTER
### SECTION: UNDERTAKING
### SECTION: DECLARATION
### SECTION: ELIGIBILITY CONFIRMATION
### SECTION: COMPLIANCE STATEMENT
### SECTION: EXPERIENCE DETAILS
### SECTION: ATTACHMENT LIST

Each section must start with:
### SECTION:

Do NOT return JSON. Return only formatted text.
"""

    # -------------------------------
    # AI CALL (Fixed)
    # -------------------------------
    @staticmethod
    def ask_ai(prompt):
        try:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You prepare tender documents."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2
            )

            # FIXED: Correct way to fetch content
            return resp.choices[0].message.content

        except Exception as e:
            return f"AI ERROR: {str(e)}"

    # -------------------------------
    # MAIN PROCESSOR
    # -------------------------------
    @staticmethod
    def process_document(uploaded_file, company, custom_prompt=""):
        bid_text = AutoDocService.read_uploaded_file(uploaded_file)

        if not bid_text.strip():
            raise Exception("Could not extract text from uploaded file.")

        prompt = AutoDocService.build_prompt(company, bid_text, custom_prompt)

        result = AutoDocService.ask_ai(prompt)

        if result.startswith("AI ERROR"):
            raise Exception(result)

        return result
